"""Extracción de texto de PDF/DOCX (spec sección 2, último parser del paso
6: "Documentos PDF/DOCX -> extracción de texto para pasada de NER").

A diferencia de csv_parser.py/chat_parser.py/server_log_parser.py, este
parser NO crea nodos de entidad por sí solo -- su único trabajo es guardar
el artefacto (hash + nodo Archivo, misma trazabilidad de siempre) y
devolver el texto extraído, LISTO para pasar a `ner.propose_entities`
(paso aparte, mismo criterio que exif_parser.describe_image_content: la
extracción de texto es determinística, la extracción de ENTIDADES desde
ese texto es responsabilidad del pipeline de NER ya existente, con su
propio estado pendiente-de-confirmación -- no se duplica esa lógica acá).

Detección de formato por firma binaria real (los primeros bytes del
archivo), no por la extensión del nombre -- un archivo mal nombrado no
tiene por qué frustrar la ingesta, y confiar en la extensión sería
exactamente el tipo de suposición no verificada que el resto del módulo
evita en todos lados."""

from __future__ import annotations

import io
from pathlib import Path

import docx
import pypdf

from . import artifact_store, case_store
from .models import make_archivo

_MIME_TYPES = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def detect_document_format(data: bytes) -> str | None:
    if data.startswith(b"%PDF"):
        return "pdf"
    if data.startswith(b"PK\x03\x04"):  # firma ZIP -- DOCX es un contenedor OOXML (ZIP) real
        return "docx"
    return None


def extract_pdf_text(data: bytes) -> str:
    reader = pypdf.PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(p for p in pages if p.strip())


def extract_docx_text(data: bytes) -> str:
    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    # el texto de tablas también es real y puede importar (ej. un informe con
    # una tabla de contactos/transacciones) -- no se descarta solo porque no
    # está en un párrafo suelto.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


_EXTRACTORS = {"pdf": extract_pdf_text, "docx": extract_docx_text}


def ingest_document(
    *, cases_dir: str | Path, keys_dir: str | Path, artifact_store_dir: str | Path, case_id: str,
    doc_bytes: bytes, original_filename: str, ingested_by: str,
) -> dict:
    doc_format = detect_document_format(doc_bytes)
    if doc_format is None:
        raise ValueError(
            f"No pude reconocer el formato de '{original_filename}' -- ni PDF ni DOCX (verificado por la "
            "firma real de los primeros bytes del archivo, no por su extensión)."
        )

    record = artifact_store.store_artifact(artifact_store_dir, doc_bytes, original_filename, ingested_by)
    archivo_node = make_archivo(
        nombre=original_filename, sha256=record.sha256, tamano=record.size,
        mime=record.mime or _MIME_TYPES[doc_format],
    )
    archivo_node = case_store.add_node(cases_dir, keys_dir, case_id, archivo_node)

    texto = _EXTRACTORS[doc_format](doc_bytes)
    return {"archivo": archivo_node, "texto": texto, "formato": doc_format}
