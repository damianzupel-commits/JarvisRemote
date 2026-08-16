"""Tests de app/investigation/doc_parser.py (paso 6, último parser:
extracción de texto de PDF/DOCX). PDF real construido a mano (PDF mínimo
válido, con tabla xref real -- no una librería de terceros nueva solo para
fixtures de test) y DOCX real construido con python-docx -- nada mockeado
en la extracción en sí."""

from __future__ import annotations

import io
from types import SimpleNamespace

import docx
import pytest

from app.investigation import case_store, doc_parser, keys, ner
from app.investigation.models import NodeType


@pytest.fixture()
def keys_dir(tmp_path):
    d = tmp_path / "keys"
    keys.ensure_keypair(d)
    return d


@pytest.fixture()
def cases_dir(tmp_path):
    d = tmp_path / "cases"
    case_store.create_case(d, "caso-1", "Caso de prueba")
    return d


@pytest.fixture()
def artifact_store_dir(tmp_path):
    return tmp_path / "artifacts"


def _make_minimal_pdf(text: str) -> bytes:
    """PDF mínimo pero real y válido -- un objeto Catalog/Pages/Page/Font/
    Contents con tabla xref de offsets reales, no un archivo truncado o
    inventado. Validado contra pypdf real antes de usarse acá."""
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 4 0 R >> >> "
        b"/MediaBox [0 0 300 144] /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    stream_content = f"BT /F1 18 Tf 10 100 Td ({text}) Tj ET".encode("latin-1")
    objects.append(
        b"<< /Length " + str(len(stream_content)).encode() + b" >>\nstream\n" + stream_content + b"\nendstream"
    )

    header = b"%PDF-1.4\n"
    body_parts = []
    offsets = []
    pos = len(header)
    for i, obj in enumerate(objects, start=1):
        offsets.append(pos)
        part = f"{i} 0 obj\n".encode() + obj + b"\nendobj\n"
        body_parts.append(part)
        pos += len(part)

    xref_offset = pos
    xref = f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode()
    for off in offsets:
        xref += f"{off:010d} 00000 n \n".encode()
    trailer = f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF".encode()

    return header + b"".join(body_parts) + xref + trailer


def _make_docx(paragraphs: list[str], table_cells: list[list[str]] | None = None) -> bytes:
    document = docx.Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table_cells:
        table = document.add_table(rows=len(table_cells), cols=len(table_cells[0]))
        for row_idx, row in enumerate(table_cells):
            for col_idx, value in enumerate(row):
                table.cell(row_idx, col_idx).text = value
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _fake_response(content: str):
    message = SimpleNamespace(content=content)
    return SimpleNamespace(choices=[SimpleNamespace(message=message)])


# --- detección de formato -----------------------------------------------------------

def test_detect_document_format_recognizes_pdf_by_signature():
    assert doc_parser.detect_document_format(_make_minimal_pdf("x")) == "pdf"


def test_detect_document_format_recognizes_docx_by_signature():
    assert doc_parser.detect_document_format(_make_docx(["x"])) == "docx"


def test_detect_document_format_returns_none_for_unrecognized_content():
    assert doc_parser.detect_document_format(b"esto no es ni PDF ni DOCX") is None


def test_detect_document_format_ignores_a_wrong_extension_and_checks_real_bytes():
    """La detección es por firma real, no por el nombre -- un archivo mal
    nombrado no tiene que frustrar la ingesta."""
    pdf_bytes = _make_minimal_pdf("contenido real")
    assert doc_parser.detect_document_format(pdf_bytes) == "pdf"  # aunque el caller lo llame "informe.docx"


# --- extracción de texto -------------------------------------------------------------

def test_extract_pdf_text_reads_the_real_text():
    pdf_bytes = _make_minimal_pdf("Hola Mundo de prueba")

    text = doc_parser.extract_pdf_text(pdf_bytes)

    assert "Hola Mundo de prueba" in text


def test_extract_docx_text_reads_paragraphs():
    docx_bytes = _make_docx(["Primer parrafo real", "Segundo parrafo real"])

    text = doc_parser.extract_docx_text(docx_bytes)

    assert "Primer parrafo real" in text
    assert "Segundo parrafo real" in text


def test_extract_docx_text_includes_table_cell_content():
    docx_bytes = _make_docx(["Informe"], table_cells=[["Juan Perez", "juan@example.com"]])

    text = doc_parser.extract_docx_text(docx_bytes)

    assert "Juan Perez" in text
    assert "juan@example.com" in text


# --- ingest_document -----------------------------------------------------------------

def test_ingest_document_creates_a_real_archivo_node_for_pdf(cases_dir, keys_dir, artifact_store_dir):
    result = doc_parser.ingest_document(
        cases_dir=cases_dir, keys_dir=keys_dir, artifact_store_dir=artifact_store_dir, case_id="caso-1",
        doc_bytes=_make_minimal_pdf("texto de prueba"), original_filename="informe.pdf", ingested_by="damian",
    )

    assert result["archivo"].tipo == NodeType.ARCHIVO
    assert result["formato"] == "pdf"
    assert "texto de prueba" in result["texto"]


def test_ingest_document_creates_a_real_archivo_node_for_docx(cases_dir, keys_dir, artifact_store_dir):
    result = doc_parser.ingest_document(
        cases_dir=cases_dir, keys_dir=keys_dir, artifact_store_dir=artifact_store_dir, case_id="caso-1",
        doc_bytes=_make_docx(["Contenido real del documento"]), original_filename="informe.docx", ingested_by="damian",
    )

    assert result["archivo"].tipo == NodeType.ARCHIVO
    assert result["formato"] == "docx"
    assert "Contenido real del documento" in result["texto"]


def test_ingest_document_raises_a_clear_error_for_unrecognized_format(cases_dir, keys_dir, artifact_store_dir):
    with pytest.raises(ValueError, match="No pude reconocer el formato"):
        doc_parser.ingest_document(
            cases_dir=cases_dir, keys_dir=keys_dir, artifact_store_dir=artifact_store_dir, case_id="caso-1",
            doc_bytes=b"no es un documento valido", original_filename="mystery.pdf", ingested_by="damian",
        )


def test_reingesting_the_same_document_resolves_to_the_same_archivo_node(cases_dir, keys_dir, artifact_store_dir):
    pdf_bytes = _make_minimal_pdf("mismo contenido")

    first = doc_parser.ingest_document(
        cases_dir=cases_dir, keys_dir=keys_dir, artifact_store_dir=artifact_store_dir, case_id="caso-1",
        doc_bytes=pdf_bytes, original_filename="informe.pdf", ingested_by="damian",
    )
    second = doc_parser.ingest_document(
        cases_dir=cases_dir, keys_dir=keys_dir, artifact_store_dir=artifact_store_dir, case_id="caso-1",
        doc_bytes=pdf_bytes, original_filename="informe.pdf", ingested_by="damian",
    )

    assert first["archivo"].id == second["archivo"].id
    archivos = [n for n in case_store.read_nodes(cases_dir, "caso-1") if n.tipo == NodeType.ARCHIVO]
    assert len(archivos) == 1


# --- integración real con el pipeline de NER ya existente -----------------------------

@pytest.mark.anyio
async def test_extracted_text_can_be_fed_into_the_existing_ner_pipeline(monkeypatch, cases_dir, keys_dir, artifact_store_dir):
    result = doc_parser.ingest_document(
        cases_dir=cases_dir, keys_dir=keys_dir, artifact_store_dir=artifact_store_dir, case_id="caso-1",
        doc_bytes=_make_minimal_pdf("Juan Perez firmo el contrato"), original_filename="contrato.pdf", ingested_by="damian",
    )

    payload = (
        '[{"tipo": "Persona", "campos": {"etiqueta": "Juan Perez", "alias": [], "confianza": 0.7}, '
        '"texto_fuente": "Juan Perez firmo el contrato", "confianza_extraccion": 0.8, "razon": "nombre propio en el documento"}]'
    )

    async def fake_create(**kwargs):
        return _fake_response(payload)

    monkeypatch.setattr(ner.client.chat.completions, "create", fake_create)
    ner_result = await ner.propose_entities(result["texto"], result["archivo"].id)
    ner.save_proposals(cases_dir, "caso-1", ner_result.proposals)

    pending = ner.read_proposals(cases_dir, "caso-1", status="pendiente")
    assert len(pending) == 1
    assert pending[0].artefacto_origen == result["archivo"].id
